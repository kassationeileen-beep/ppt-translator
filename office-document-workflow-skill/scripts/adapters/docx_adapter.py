"""DOCX extraction/writeback adapter using python-docx."""
from __future__ import annotations

from pathlib import Path
from typing import Dict, Iterable, Iterator, List, Tuple

from docx import Document

from scripts.core.document_model import RiskFlags, StyleMetadata, TextBlock


def _color_to_hex(run) -> str | None:
    color = run.font.color
    if color is None or color.rgb is None:
        return None
    return str(color.rgb)


def _font_size_points(run) -> float | None:
    return float(run.font.size.pt) if run.font.size is not None else None


class DocxAdapter:
    document_type = "docx"

    def __init__(self, path: str | Path):
        self.path = Path(path)
        self.document = Document(str(self.path))

    def extract_text_blocks(self) -> List[TextBlock]:
        blocks: List[TextBlock] = []
        for section_path, paragraph in self._iter_paragraphs():
            paragraph_index = int(section_path.rsplit("/", 1)[-1].replace("p", "")) if "/p" in section_path else 0
            para_style = paragraph.style.name if paragraph.style is not None else None
            alignment = str(paragraph.alignment) if paragraph.alignment is not None else None
            for run_index, run in enumerate(paragraph.runs):
                if not run.text:
                    continue
                block_id = f"docx:{section_path}:r{run_index}"
                blocks.append(TextBlock(
                    document_type="docx",
                    block_id=block_id,
                    paragraph_index=paragraph_index,
                    run_index=run_index,
                    section_path=section_path,
                    original_text=run.text,
                    style=StyleMetadata(
                        font_name=run.font.name,
                        font_size=_font_size_points(run),
                        bold=run.bold,
                        italic=run.italic,
                        underline=run.underline,
                        color=_color_to_hex(run),
                        alignment=alignment,
                        paragraph_style=para_style,
                    ),
                    risks=RiskFlags(is_table_text=section_path.startswith("body/table")),
                    context={"kind": self._kind_from_path(section_path)},
                ))
            if not paragraph.runs and paragraph.text:
                blocks.append(TextBlock(
                    document_type="docx",
                    block_id=f"docx:{section_path}:text",
                    paragraph_index=paragraph_index,
                    section_path=section_path,
                    original_text=paragraph.text,
                    style=StyleMetadata(alignment=alignment, paragraph_style=para_style),
                    context={"kind": self._kind_from_path(section_path), "paragraph_level": True},
                ))
        return blocks

    def write_translations(self, blocks: Iterable[TextBlock], output_path: str | Path) -> None:
        block_map: Dict[str, TextBlock] = {block.block_id: block for block in blocks if block.translated_text and not block.risks.requires_manual_review}
        for section_path, paragraph in self._iter_paragraphs():
            if paragraph.runs:
                for run_index, run in enumerate(paragraph.runs):
                    block = block_map.get(f"docx:{section_path}:r{run_index}")
                    if block:
                        run.text = block.translated_text or run.text
            else:
                block = block_map.get(f"docx:{section_path}:text")
                if block:
                    paragraph.text = block.translated_text or paragraph.text
        self.document.save(str(output_path))

    def _iter_paragraphs(self) -> Iterator[Tuple[str, object]]:
        for index, paragraph in enumerate(self.document.paragraphs):
            yield f"body/p{index}", paragraph
        for table_index, table in enumerate(self.document.tables):
            yield from self._iter_table(table, f"body/table{table_index}")
        for section_index, section in enumerate(self.document.sections):
            for index, paragraph in enumerate(section.header.paragraphs):
                yield f"section{section_index}/header/p{index}", paragraph
            for table_index, table in enumerate(section.header.tables):
                yield from self._iter_table(table, f"section{section_index}/header/table{table_index}")
            for index, paragraph in enumerate(section.footer.paragraphs):
                yield f"section{section_index}/footer/p{index}", paragraph
            for table_index, table in enumerate(section.footer.tables):
                yield from self._iter_table(table, f"section{section_index}/footer/table{table_index}")

    def _iter_table(self, table, base_path: str) -> Iterator[Tuple[str, object]]:
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    yield f"{base_path}/r{row_index}c{col_index}/p{paragraph_index}", paragraph
                for nested_index, nested in enumerate(cell.tables):
                    yield from self._iter_table(nested, f"{base_path}/r{row_index}c{col_index}/table{nested_index}")

    @staticmethod
    def _kind_from_path(section_path: str) -> str:
        if "header" in section_path:
            return "header"
        if "footer" in section_path:
            return "footer"
        if "table" in section_path:
            return "table"
        return "paragraph"
